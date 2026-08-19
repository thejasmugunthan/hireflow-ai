import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_docx(filename="HireFlow_AI_Use_Case_Documentation.docx"):
    doc = docx.Document()

    # Set page margins (0.75 in)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # ── TITLE & SUBTITLE ──────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("HireFlow AI")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(10, 102, 194) # #0A66C2

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Next-Generation AI-Powered Applicant Tracking & Talent Evaluation Platform\nSystem Use Case Specification & Architecture Documentation")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    # ── METADATA TABLE ────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=3, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        [("Document Version:", True), ("1.0.0 (Production Release)", False), ("Target Platform:", True), ("AWS Amplify (Web) + Render (API)", False)],
        [("Author / Lead:", True), ("Thejas", False), ("Database:", True), ("MongoDB Atlas Cloud Cluster", False)],
        [("Date of Release:", True), ("August 2026", False), ("AI Engine:", True), ("OpenAI GPT-4o-mini + Word NLP", False)],
    ]

    for r_idx, row in enumerate(meta_data):
        for c_idx, (text, is_bold) in enumerate(row):
            cell = meta_table.cell(r_idx, c_idx)
            cell.text = text
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.bold = is_bold
            if is_bold:
                p.runs[0].font.color.rgb = RGBColor(15, 23, 42)
            else:
                p.runs[0].font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph() # Spacer

    # ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────
    h1 = doc.add_heading(level=1)
    r = h1.add_run("1. Executive Summary & Solution Overview")
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph("HireFlow AI is a full-stack, enterprise-grade Talent Operating System and Applicant Tracking System (ATS) engineered to bridge candidate discovery with autonomous AI-driven candidate evaluation. It replaces fragmented hiring spreadsheets with automated resume decoding, authenticity/plagiarism detection, strict forward-only hiring stage governance, and cloud-native PDF persistence.")
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(10.5)

    # Architecture Table
    arch_table = doc.add_table(rows=6, cols=3)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_headers = ["System Layer", "Technology Stack", "Core Responsibilities"]
    for c_idx, h in enumerate(arch_headers):
        cell = arch_table.cell(0, c_idx)
        cell.text = h
        set_cell_background(cell, "EAF4FF")
        p = cell.paragraphs[0]
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(10, 102, 194)

    arch_rows = [
        ("Frontend UI/UX", "React 18, Vite, Tailwind CSS, Lucide Icons", "Responsive public job portal, dedicated mobile candidate cards, recruiter command center."),
        ("Cloud Hosting", "AWS Amplify (Web) + Render (Node.js API)", "Automated CI/CD git-triggered builds, SSL termination, and global CDN delivery."),
        ("Backend Server", "Node.js, Express, Multer, JWT, CORS", "REST API endpoints, JWT security guard, resume binary streaming, stage transitions."),
        ("Database Layer", "MongoDB Atlas (Mongoose ODM)", "Cloud Base64 PDF storage, candidate indexing, interview logs, duplicate protection."),
        ("Intelligence Engine", "OpenAI GPT-4o-mini & NLP Heuristic", "Skill extraction, role match scoring (0-100%), plagiarism risk & originality audit."),
    ]

    for r_idx, (layer, tech, resp) in enumerate(arch_rows, start=1):
        for c_idx, val in enumerate([layer, tech, resp]):
            cell = arch_table.cell(r_idx, c_idx)
            cell.text = val
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                p.runs[0].font.bold = True

    doc.add_paragraph()

    # ── 2. PRIMARY SYSTEM ACTORS ──────────────────────────────────────────
    h1 = doc.add_heading(level=1)
    r = h1.add_run("2. Primary System Actors")
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(15, 23, 42)

    actors_table = doc.add_table(rows=4, cols=2)
    actors_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, h in enumerate(["Actor", "Description & Operational Scope"]):
        cell = actors_table.cell(0, c_idx)
        cell.text = h
        set_cell_background(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

    actor_rows = [
        ("Job Applicant (Candidate)", "Public user browsing active positions, submitting PDF resumes with personal credentials, and receiving fast-track application confirmations."),
        ("Recruiter / Hiring Team", "Authenticated admin managing requisitions, inspecting AI candidate scoring, viewing streamed PDF resumes, conducting multi-round interview ratings, and advancing candidates along strict stage gates."),
        ("AI Intelligence Service", "Autonomous background engine that extracts keywords, assesses technical depth, calculates originality vs plagiarism risk, and compiles strengths and gaps."),
    ]

    for r_idx, (act, desc) in enumerate(actor_rows, start=1):
        for c_idx, val in enumerate([act, desc]):
            cell = actors_table.cell(r_idx, c_idx)
            cell.text = val
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                p.runs[0].font.bold = True

    doc.add_paragraph()

    # ── 3. DETAILED USE CASES ─────────────────────────────────────────────
    h1 = doc.add_heading(level=1)
    r = h1.add_run("3. Detailed Use Case Specifications")
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(15, 23, 42)

    use_cases = [
        {
            "id": "UC-01", "title": "Job Requisition Discovery & Filter Search",
            "actor": "Job Applicant / Public Visitor",
            "precond": "Public access to HireFlow AI landing page.",
            "trigger": "Candidate visits home URL or clicks 'Open Positions'.",
            "flow": "1. System loads active job listings from database via GET /api/jobs.\n2. Candidate searches by role keywords, skills, or location.\n3. Candidate filters by employment type (Full-Time, Contract, Internship).\n4. Candidate clicks 'Apply Now' on target requisition.",
            "alt": "4a. If no jobs match search criteria, system presents helpful empty state with reset filters button.",
            "post": "Candidate is navigated to Apply page with target jobId pre-selected."
        },
        {
            "id": "UC-02", "title": "Application Submission & Atlas PDF Storage",
            "actor": "Job Applicant",
            "precond": "Target job is active and accepting applications.",
            "trigger": "Candidate submits application form with PDF resume upload.",
            "flow": "1. Candidate inputs Full Name, Email, Phone, LinkedIn, and optional cover note.\n2. Candidate uploads PDF resume file (max 5MB).\n3. Backend parses multipart form data via Multer and converts PDF buffer to Base64 string.\n4. System saves Candidate document and Application record in MongoDB Atlas.\n5. System initializes hiring stage at Applied and triggers asynchronous AI parsing.\n6. UI redirects candidate to Application Success page with summary details.",
            "alt": "3a. Duplicate email applied to same job: System returns 409 Conflict with clear error message.\n3b. Invalid file format / size > 5MB: System rejects upload with 400 Bad Request.",
            "post": "Application and PDF resume are permanently persisted in MongoDB Atlas cloud."
        },
        {
            "id": "UC-03", "title": "AI Resume Extraction & Skill Alignment",
            "actor": "AI Intelligence Engine / Recruiter",
            "precond": "Application submitted with PDF resume data.",
            "trigger": "Automated on application creation or triggered via 'Re-run Analysis'.",
            "flow": "1. AI Service decodes Base64 resume buffer and extracts text.\n2. OpenAI GPT-4o-mini prompt analyzes resume against target job description.\n3. System extracts candidate skills list and determines matched vs missing competencies.\n4. Computes Match Score (0-100%), verified strengths list, and potential interview gap probes.\n5. Results saved to application document under aiAnalysis schema.",
            "alt": "2a. OpenAI API key unavailable or rate limited: System seamlessly falls back to local NLP heuristic word-boundary parser without user disruption.",
            "post": "Candidate profile displays accurate Match Score radial ring and extracted skills breakdown."
        },
        {
            "id": "UC-04", "title": "Resume Plagiarism & Authenticity Audit",
            "actor": "AI Intelligence Engine / Recruiter",
            "precond": "Resume text extracted and ready for evaluation.",
            "trigger": "Triggered as part of AI Candidate Insights pipeline.",
            "flow": "1. System scans resume text for boilerplate buzzword density, template clichés, and project specificity.\n2. Computes Originality Score (%) and Plagiarism Risk (%).\n3. Assigns risk categorization: Low Risk (<=20%), Standard Phrasing (21-40%), High Template Plagiarism (>40%).\n4. Renders dual-color progress gauge and authenticity observation flag on recruiter card.",
            "alt": "1a. Short or non-technical resume: System flags lack of quantified metrics and raises alert.",
            "post": "Recruiter is empowered with objective authenticity rating to screen out copied resumes."
        },
        {
            "id": "UC-05", "title": "Strict Forward-Only Stage Governance",
            "actor": "Recruiter / Hiring Manager",
            "precond": "Authenticated admin viewing candidate details at /admin/applications/:id.",
            "trigger": "Recruiter clicks 'Update Stage' button.",
            "flow": "1. System opens Stage Decision Modal displaying current stage (e.g. Applied).\n2. Modal presents two forward paths: Selected / Advance (-> R1) OR Rejected (-> Reject).\n3. Recruiter enters mandatory decision rationale note.\n4. Backend validates transition against strict forward state machine:\n   • Applied -> R1 | Reject\n   • R1 -> R2 | R1 Reject\n   • R2 -> R3 | R2 Reject\n   • R3 -> Approved | R3 Reject\n5. Stage history log updated with timestamp and decision note.",
            "alt": "4a. Recruiter attempts illegal backward move: System rejects transition with validation error.\n4b. Candidate already in terminal stage (Approved / Reject): Modal locks stage selector.",
            "post": "Candidate stage updated and reflected across dashboard, timeline, and mobile cards."
        },
        {
            "id": "UC-06", "title": "Inline Cloud PDF Resume Streaming",
            "actor": "Recruiter / Hiring Manager",
            "precond": "Candidate document contains Base64 resume buffer in Atlas.",
            "trigger": "Recruiter clicks 'View Resume (PDF)' in Admin portal.",
            "flow": "1. Request dispatched to GET /api/applications/:id/resume.\n2. Backend retrieves candidate.resumeData from MongoDB Atlas.\n3. Decodes Base64 buffer to binary Buffer.from(data, 'base64').\n4. Sets response headers: Content-Type: application/pdf, Content-Disposition: inline.\n5. Browser opens resume PDF in a dedicated tab with native PDF controls (zoom, print, save).",
            "alt": "2a. Resume data missing / legacy upload: Backend attempts local file fallback or returns 404.",
            "post": "Recruiter inspects original candidate PDF with zero dependency on ephemeral cloud storage."
        },
        {
            "id": "UC-07", "title": "Structured Interview Feedback & Ratings",
            "actor": "Recruiter / Technical Interviewer",
            "precond": "Candidate is in active interview round (R1, R2, or R3).",
            "trigger": "Interviewer logs feedback in 'Interview Notes & Evaluation' section.",
            "flow": "1. Recruiter selects Interview Round (R1 Technical, R2 Deep Tech, R3 System Design, HR).\n2. Assigns 1 to 5 Star Rating.\n3. Enters detailed technical observations, questions asked, and strengths.\n4. Clicks 'Save Interview Note'.\n5. System persists note in candidate interviewNotes array and refreshes timeline.",
            "alt": "3a. Empty notes field: System validates input and displays inline prompt.",
            "post": "Feedback is permanently chronologically indexed for collaborative hiring decisions."
        }
    ]

    for uc in use_cases:
        t = doc.add_table(rows=7, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        cell_header = t.cell(0, 0)
        cell_header.merge(t.cell(0, 1))
        cell_header.text = f"{uc['id']}: {uc['title']}"
        set_cell_background(cell_header, "0A66C2")
        set_cell_margins(cell_header, top=100, bottom=100, left=150, right=150)
        p = cell_header.paragraphs[0]
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

        fields = [
            ("Primary Actor:", uc['actor']),
            ("Preconditions:", uc['precond']),
            ("Trigger:", uc['trigger']),
            ("Main Success Flow:", uc['flow']),
            ("Exceptions / Alternates:", uc['alt']),
            ("Postconditions:", uc['post']),
        ]

        for idx, (label, val) in enumerate(fields, start=1):
            c_lbl = t.cell(idx, 0)
            c_lbl.text = label
            set_cell_background(c_lbl, "F8FAFC")
            set_cell_margins(c_lbl, top=80, bottom=80, left=120, right=120)
            p = c_lbl.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(15, 23, 42)

            c_val = t.cell(idx, 1)
            c_val.text = val
            set_cell_background(c_val, "FFFFFF")
            set_cell_margins(c_val, top=80, bottom=80, left=120, right=120)
            p = c_val.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(51, 65, 85)

        doc.add_paragraph() # Spacer

    # ── 4. SECURITY & GOVERNANCE MATRIX ───────────────────────────────────
    h1 = doc.add_heading(level=1)
    r = h1.add_run("4. System Governance, Security & Quality Matrix")
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(15, 23, 42)

    sec_table = doc.add_table(rows=6, cols=3)
    sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, h in enumerate(["Governance Dimension", "Implementation Mechanism", "Impact & Guarantee"]):
        cell = sec_table.cell(0, c_idx)
        cell.text = h
        set_cell_background(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

    sec_rows = [
        ("Authentication & RBAC", "JWT Bearer Tokens + bcrypt salted password hashing", "Guarantees admin endpoints are shielded from unauthorized access."),
        ("Duplicate Guard", "MongoDB compound index on (jobId + candidate email)", "Prevents candidates from spamming duplicate applications for same role."),
        ("Pipeline Integrity", "Deterministic finite state machine (FSM) validation", "Enforces strict forward-only progression; prevents illegal stage regressions."),
        ("Ephemeral Cloud Resilience", "Atlas Base64 direct buffer storage & streaming", "Ensures 100% resume persistence across container restarts on Render/Amplify."),
        ("Mobile Responsiveness", "Dedicated mobile cards view (sm:hidden / md:table)", "Flawless touch usability on 360px+ devices without layout clipping."),
    ]

    for r_idx, (dim, mech, imp) in enumerate(sec_rows, start=1):
        for c_idx, val in enumerate([dim, mech, imp]):
            cell = sec_table.cell(r_idx, c_idx)
            cell.text = val
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                p.runs[0].font.bold = True

    doc.add_paragraph()

    # ── 5. CONCLUSION ─────────────────────────────────────────────────────
    h1 = doc.add_heading(level=1)
    r = h1.add_run("5. Document Authorization & Sign-Off")
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph("This Use Case Document certifies that the HireFlow AI application meets all architectural, functional, and governance specifications for production deployment across AWS Amplify and Render cloud environments.")
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(10)

    sign_table = doc.add_table(rows=2, cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_data = [
        [("Prepared By:", True), ("Thejas (Lead Engineer)", False), ("Project:", True), ("HireFlow AI Platform", False)],
        [("Status:", True), ("Verified & Production Ready", False), ("Approved By:", True), ("System Architect", False)],
    ]
    for r_idx, row in enumerate(sign_data):
        for c_idx, (text, is_bold) in enumerate(row):
            cell = sign_table.cell(r_idx, c_idx)
            cell.text = text
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.bold = is_bold

    doc.save(filename)
    print(f"DOCX Successfully Generated: {filename}")

if __name__ == "__main__":
    create_docx()
